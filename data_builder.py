import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_rhp_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('RED HERRING PROSPECTUS', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Dated December 10, 2025\nPlease read section 32 of the Companies Act, 2013\n100% Book Built Offer').bold = True

    # Company Details
    doc.add_heading('KSH INTERNATIONAL LIMITED', level=1)
    doc.add_paragraph('Corporate Identity Number: U28129PN1979PLC141032')
    doc.add_paragraph('Formerly known as Bhandary Metal Extrusion Private Limited')

    # Registered & Corporate Office
    p_reg = doc.add_paragraph()
    p_reg.add_run('Registered Office: ').bold = True
    p_reg.add_run('11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed, Pune – 410 501, Maharashtra, India\n')
    p_reg.add_run('Corporate Office: ').bold = True
    p_reg.add_run('201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune – 411 045, Maharashtra, India\n')
    p_reg.add_run('Contact Person: ').bold = True
    p_reg.add_run('Sarthak Malvadkar, Company Secretary and Compliance Officer\n')
    p_reg.add_run('Telephone: ').bold = True
    p_reg.add_run('+91 20 45053237\n')
    p_reg.add_run('Email: ').bold = True
    p_reg.add_run('cs.connect@kshinternational.com\n')
    p_reg.add_run('Website: ').bold = True
    p_reg.add_run('www.kshinternational.com')

    # Promoters
    doc.add_heading('OUR PROMOTERS', level=2)
    doc.add_paragraph(
        'KUSHAL SUBBAYYA HEGDE, PUSHPA KUSHAL HEGDE, RAJESH KUSHAL HEGDE, ROHIT KUSHAL HEGDE, '
        'RAKHI GIRIJA SHETTY, DHAULAGIRI FAMILY TRUST, EVEREST FAMILY TRUST, MAKALU FAMILY TRUST, '
        'BROAD FAMILY TRUST, ANNAPURNA FAMILY TRUST, KANCHENJUNGA FAMILY TRUST AND WATERLOO INDUSTRIAL PARK VI PRIVATE LIMITED'
    )

    # Lead Managers & Registrar Table
    doc.add_heading('BOOK RUNNING LEAD MANAGERS & REGISTRAR', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Entity Name'
    hdr_cells[1].text = 'Contact Person'
    hdr_cells[2].text = 'Contact Info'

    brm_data = [
        ('Nuvama Wealth Management Limited', 'Lokesh Shah / Soumavo Sarkar / Prakash Boricha', 'Email: ksh.ipo@nuvama.com, Tel: +91 22 4009 4400'),
        ('ICICI Securities Limited', 'Kishan Rastogi / Abhijit Diwan', 'Email: ksh@icicisecurities.com, Tel: +91 22 6807 7100'),
        ('MUFG Intime India Private Limited (Registrar)', 'Shanti Gopalkrishnan', 'Email: kshinternational.ipo@in.mpms.mufg.com, Tel: +91 81081 14949')
    ]

    for entity, contact, info in brm_data:
        row_cells = table.add_row().cells
        row_cells[0].text = entity
        row_cells[1].text = contact
        row_cells[2].text = info

    # Board of Directors Table
    doc.add_heading('BOARD OF DIRECTORS', level=2)
    dir_table = doc.add_table(rows=1, cols=4)
    dir_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    d_hdr = dir_table.rows[0].cells
    d_hdr[0].text = 'Name'
    d_hdr[1].text = 'Designation'
    d_hdr[2].text = 'DIN'
    d_hdr[3].text = 'Address'

    directors = [
        ('Kushal Subbayya Hegde', 'Chairman & Executive Director', '00135070', 'S. no. 245/ 104, Pushpakamal, Deccan Gymkhana Society, lane no. 3 Prabhat Road, Pune – 411 004, Maharashtra, India'),
        ('Rajesh Kushal Hegde', 'Managing Director', '00114193', '12 Buena Monte, NCL co-operative housing society, Panchvati, Pashan, Pune – 411 008, Maharashtra, India'),
        ('Rohit Kushal Hegde', 'Joint Managing Director', '00134926', 'Pushpakamal Apartment, Flat – 1, S. no. 245/ 104, Prabhat Road, Lane no. 3, Shivaji Nagar, Pune – 411 004, Maharashtra, India'),
        ('Rakhi Girija Shetty', 'Whole-time Director', '03124510', 'S. no. 245/ 104, Pushpakamal, Deccan Gymkhana Society, lane no. 3 Prabhat Road, Erandawane, Pune – 411 004, Maharashtra, India'),
        ('Dinesh Hirachand Munot', 'Independent Director', '00049801', 'Pratik Bunglow, Senapati Bapat Road, behind Sahara Hotel, Shivajinagar, Model Colony, Pune – 411 016, Maharashtra, India'),
        ('Ajay Shriram Patil', 'Independent Director', '01217000', '602, Gopalkrupa Apartment, Bhonde colony, Prabhat Road, Erandawane, Pune – 411 004, Maharashtra, India'),
        ('Ram Kumar Tiwari', 'Independent Director', '10938958', 'A-259, JK Road, Minal Residency, Huzur, Govindpura, Bhopal – 462 023, Madhya Pradesh, India'),
        ('Indu Jacob', 'Independent Director', '05293084', 'A29, Abhimanshree Society, Pashan Road, Pune – 411 008, Maharashtra, India')
    ]

    for dname, desig, din, addr in directors:
        r_cells = dir_table.add_row().cells
        r_cells[0].text = dname
        r_cells[1].text = desig
        r_cells[2].text = din
        r_cells[3].text = addr

    # Other Key Executives & Auditors
    doc.add_heading('KEY MANAGERIAL PERSONNEL & AUDITORS', level=2)
    p_kmp = doc.add_paragraph()
    p_kmp.add_run('Chief Executive Officer (CEO): ').bold = True
    p_kmp.add_run('Sandesh Bhagwat\n')
    p_kmp.add_run('Chief Financial Officer (CFO): ').bold = True
    p_kmp.add_run('Amod Joshi\n')
    p_kmp.add_run('Independent Chartered Engineer: ').bold = True
    p_kmp.add_run('Lalit Muljibhai Sarvaiya (Registration No: M-140388)\n')
    p_kmp.add_run('Statutory Auditors: ').bold = True
    p_kmp.add_run('Kirtane & Pandit LLP, 5th Floor, Wing A, Gopal House, S. No. 127/1B/1, Karve Road, Pune – 411 038. Email: parag.pansare@kirtanepandit.com, Tel: +91 20 6729 5100\n')
    p_kmp.add_run('Legal Counsel: ').bold = True
    p_kmp.add_run('Trilegal, 10th Floor, Tower 2A & 2B, Senapati Bapat Marg, Lower Parel, Mumbai – 400 013. Email: ipo@trilegal.com, Tel: +91 22 4079 1000')

    # Banking Contacts
    doc.add_heading('BANKING CONTACTS', level=2)
    p_bank = doc.add_paragraph()
    p_bank.add_run('Citibank N.A. Contact: ').bold = True
    p_bank.add_run('Hitesh Ramani, Tel: +91 20 6606 4494, Email: hitesh.ramani@citi.com\n')
    p_bank.add_run('Export-Import Bank of India Contact: ').bold = True
    p_bank.add_run('Chitra Raste, Tel: +91 20 2640 3100, Email: pro@eximbankindia.in\n')
    p_bank.add_run('IndusInd Bank Limited Contact: ').bold = True
    p_bank.add_run('Sharmila Joshi, Tel: +91 20 26234000, Email: sharmila.joshi@indusind.com\n')
    p_bank.add_run('ICICI Bank Limited Contact: ').bold = True
    p_bank.add_run('Cherag Gyara, Tel: +91 8879770456, Email: cherag.gyara@icicibank.com\n')
    p_bank.add_run('HDFC Bank Limited Contact: ').bold = True
    p_bank.add_run('Manisha Shukla, Tel: +91 20 6769 4648, Email: manisha.shukla@hdfcbank.com / Eric Bacha / Sachin Gawade / Pravin Teli / Siddharth Jadhav / Tushar Gavankar (siddharth.jadhav@hdfcbank.com)')

    doc.add_heading('ATTACHMENT: IDENTIFICATION DOCUMENT RECORD', level=1)
    p_pan = doc.add_paragraph()
    p_pan.add_run('INCOME TAX DEPARTMENT - GOVT. OF INDIA\n').bold = True
    p_pan.add_run('Permanent Account Number (PAN): ').bold = True
    p_pan.add_run('NBWPS1951N\n')
    p_pan.add_run('Full Name: ').bold = True
    p_pan.add_run('VISHAL SINGH\n')
    p_pan.add_run('Father\'s Name: ').bold = True
    p_pan.add_run('SUGRIV SINGH\n')
    p_pan.add_run('Date of Birth: ').bold = True
    p_pan.add_run('06/05/2000\n')
    p_pan.add_run('Issuing Authority Address: ').bold = True
    p_pan.add_run('Income Tax PAN Services Unit, NSDL, 4th Floor, Mantri Sterling, Plot No. 341, Survey No. 997/8, Model Colony, Near Deep Bungalow Chowk, Pune - 411 016.\n')
    p_pan.add_run('Contact Tel: ').bold = True
    p_pan.add_run('91-20-2721 8080, Fax: 91-20-2721 8081\n')
    p_pan.add_run('Email: ').bold = True
    p_pan.add_run('tininfo@nsdl.co.in')

    doc.save('data/red_herring_prospectus.docx')
    print('Created data/red_herring_prospectus.docx successfully.')

def create_sample_ticket_logs():
    logs = """[2026-08-10 09:15:22] TICKET #10492 - Customer Query
User: Rashi Patil
Email: rashhi.patil@gmail.com
Phone: +91 9876543210
Address: 45 Park Avenue, Block C, Bandra West, Mumbai 400050, India
DOB: 14/08/1992
SSN / Govt ID: 324-55-9102
Credit Card: 4532-8910-4421-9018
IP: 192.168.1.104
Company: Acme Solutions Pvt Ltd
Issue: Billing discrepancy on invoice #INV-9921. Replaced contact info requested: John Doe (john.doe@example.com, +91 1234567645).

[2026-08-10 10:42:11] TICKET #10493 - Technical Support
User: Rohan Dey
Email: rohan.dey@gmail.com
Phone: +91 9988776655
Address: Flat 302, Green Valley Apartments, Off MG Road, Bengaluru 560001, India
DOB: 23/11/1988
SSN / Govt ID: 419-88-2301
Credit Card: 5412-7511-9023-4156
IP: 10.0.4.52
Company: TechCorp Innovations LLC
Issue: Unable to reset password. Alternate contact: Peter Parker (peter.parker@example.com).

[2026-08-10 11:30:05] TICKET #10494 - Account Escalation
User: Sarthak Malvadkar
Email: cs.connect@kshinternational.com
Phone: +91 20 45053237
Address: 201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune 411045, Maharashtra, India
DOB: 12/05/1985
Govt ID (PAN): NBWPS1951N
DIN: 00135070
IP: 172.16.254.1
Company: KSH INTERNATIONAL LIMITED
Issue: Corporate filing update for promoter Kushal Subbayya Hegde (kushal.hegde@kshinternational.com).
"""
    with open('data/sample_ticket_logs.txt', 'w', encoding='utf-8') as f:
        f.write(logs)
    print('Created data/sample_ticket_logs.txt successfully.')

if __name__ == '__main__':
    create_rhp_docx()
    create_sample_ticket_logs()
