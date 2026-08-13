import docx
from docx import Document
from typing import Dict, Any, List
from src.detector import PIIDetector
from src.redactor import PIIRedactor

class DocxRedactor:
    def __init__(self, detector: PIIDetector = None, redactor: PIIRedactor = None):
        self.detector = detector or PIIDetector()
        self.redactor = redactor or PIIRedactor(mode="synthetic")

    def _redact_paragraph(self, p) -> int:
        full_text = p.text
        if not full_text.strip():
            return 0

        entities = self.detector.detect(full_text)
        if not entities:
            return 0

        redacted_text, changes = self.redactor.redact_text(full_text, entities)
        
        # Update paragraph text while preserving style
        if p.runs:
            # Simple & robust run-preservation strategy: set first run text to redacted_text, clear others
            p.runs[0].text = redacted_text
            for run in p.runs[1:]:
                run.text = ""
        else:
            p.text = redacted_text

        return len(changes)

    def redact_file(self, input_docx_path: str, output_docx_path: str) -> Dict[str, Any]:
        doc = Document(input_docx_path)
        total_redactions = 0

        # 1. Process Body Paragraphs
        for p in doc.paragraphs:
            total_redactions += self._redact_paragraph(p)

        # 2. Process Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        total_redactions += self._redact_paragraph(p)

        # 3. Process Headers and Footers
        for section in doc.sections:
            for p in section.header.paragraphs:
                total_redactions += self._redact_paragraph(p)
            for p in section.footer.paragraphs:
                total_redactions += self._redact_paragraph(p)

        doc.save(output_docx_path)
        return {
            "input_file": input_docx_path,
            "output_file": output_docx_path,
            "total_redactions": total_redactions,
            "mappings": self.redactor.entity_map
        }
